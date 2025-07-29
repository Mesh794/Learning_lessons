import uuid
import json
import streamlit as st
from docx import Document
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_core.documents import Document as LCDocument
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_pinecone import PineconeVectorStore
from pinecone import Pinecone, ServerlessSpec


# Підключеня до Gemini і Pinecone
gemini = st.secrets.get("GEMINI_API_KEY")
pinecone_key = st.secrets.get("PINECONE_API_KEY")

embeddings = GoogleGenerativeAIEmbeddings(
    model="models/text-embedding-004",
    google_api_key=gemini
)

pc = Pinecone(api_key=pinecone_key)
index_name = "exam-itstep"
if not pc.has_index(index_name):
    pc.create_index(
        name=index_name,
        dimension=768,
        metric="cosine",
        spec=ServerlessSpec(
            cloud="aws",
            region="us-east-1"
        )
    )

index = pc.Index(index_name)
vector_store = PineconeVectorStore(index=index, embedding=embeddings)


# Обробка .docx файлу - просто with open as не працює :)
word_name = "For wokers.docx"
docx = Document(word_name)
word_chunks = []
current_title = "Вступ. Загальні положення"
current_word_text = ""

for paragraph in docx.paragraphs:
    if paragraph.style.name.startswith("Heading") and paragraph.text.strip():
        if current_word_text:
            word_chunks.append((current_title, current_word_text.strip()))
        current_title = paragraph.text.strip()
        current_word_text = ""
    else:
        current_word_text += paragraph.text.strip() + " "
if current_word_text:
    word_chunks.append((current_title, current_word_text.strip()))


# Обробка .pdf файлу - просто with open as не працює :)
pdf_name = "General.pdf"
loader = PyMuPDFLoader(pdf_name)
pdf = loader.load()
pdf_chunks = []
current_section = "1"
current_pdf_text = ""

for doc in pdf:
    lines = doc.page_content.splitlines()
    for line in lines:
        line = line.strip()
        if len(line) >= 2 and line[0].isdigit() and line[1] == ".":
            section_digit = line[0]
            if section_digit != current_section:
                if current_pdf_text:
                    pdf_chunks.append((f"Розділ {current_section}", current_pdf_text.strip()))
                current_section = section_digit
                current_text = line + " "
            else:
                current_pdf_text += line + " "
        else:
            current_pdf_text += line + " "
if current_pdf_text:
    pdf_chunks.append((f"Розділ {current_section}", current_pdf_text.strip()))


# Завантаження до Pinecone
documents = []
id_mapping = []

for source_file, chunks in [(word_name, word_chunks), (pdf_name, pdf_chunks)]:
    for section, text in chunks:
        doc_id = str(uuid.uuid4())
        metadata = {
            "source": source_file,
            "section": section,
        }
        documents.append(LCDocument(page_content=text, metadata=metadata))
        id_mapping.append({"id": doc_id, "source": source_file, "section": section})

ids = [item["id"] for item in id_mapping]
vector_store.add_documents(documents=documents, ids=ids)


# Завантаження до.json
with open("pinecone_inserts.json", "w", encoding="utf-8") as f:
    json.dump(id_mapping, f, ensure_ascii=False, indent=2)
